import io
import pandas as pd
import streamlit as st
from docx import Document
import csv

def build_mail_merge_dataframe(records: list[dict]) -> pd.DataFrame:
    rows = []
    for r in records:
        first_name = r["name"].split()[0]

        rows.append(
            {
                "FirstName": first_name,
                "Email": r["email"],
                # KEY CHANGE HERE 👇
                "MissedHomeworks": "\n".join(r["missed"]),
            }
        )

    return pd.DataFrame(rows)

def generate_missed_homework_report(
    df: pd.DataFrame,
    first_name_col: str = "First Name",
    last_name_col: str = "Last Name",
    group_col: str = "Group",
    email_col: str = "Email",
    scores_start_index: int = 4,
) -> list[dict]:
    """
    Returns a list of records, one per student who missed >= 1 assignment:
      [
        {
          "name": "First Last",
          "group": 1,
          "email": "x@y.edu",
          "missed": ["HW1", "HW3", ...]
        },
        ...
      ]

    Assumes assignment score columns start at df.columns[scores_start_index:].
    """

    required = {first_name_col, last_name_col, group_col, email_col}
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Missing required column(s): {', '.join(sorted(missing))}")

    score_columns = df.columns[scores_start_index:]
    if len(score_columns) == 0:
        raise ValueError(
            f"No assignment columns detected. Expected assignment columns from column index {scores_start_index} onward."
        )

    results: list[dict] = []

    for _, row in df.iterrows():
        first = str(row[first_name_col]).strip()
        last = str(row[last_name_col]).strip()
        name = f"{first} {last}".strip()

        group_val = row[group_col]
        email_val = str(row[email_col]).strip()

        missed = []
        for col in score_columns:
            val = row[col]
            # Only treat numeric 0 as missed (same as your original behavior).
            # If you want blanks (NaN) to count as missed too, I can change it.
            if pd.notna(val) and val == 0:
                missed.append(str(col))

        # Only keep students with at least one missed assignment
        if missed:
            results.append(
                {
                    "name": name,
                    "group": group_val,
                    "email": email_val,
                    "missed": missed,
                }
            )

    return results


def build_word_report_bytes(records: list[dict], title: str) -> bytes:
    doc = Document()
    doc.add_heading(title, 0)

    if not records:
        doc.add_paragraph("No students with missed homework were found.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    for r in records:
        # Example: Damilola Popoola, 1, dami@iastate.edu
        header_line = f"{r['name']}, {r['group']}, {r['email']}"
        doc.add_heading(header_line, level=1)

        for a in r["missed"]:
            doc.add_paragraph(f" {a}", style="ListBullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main():
    st.set_page_config(page_title="Missed Homework Report Generator v1", layout="centered")
    st.title("Missed Homework Report Generator v1")
    st.caption("Upload an Excel file and download a Word report listing only students with score = 0 on any assignment.")

    with st.sidebar:
        st.header("Options")
        title = st.text_input("Report title", value="Missed Homework Report")

        st.subheader("Column names (must match Excel headers)")
        first_name_col = st.text_input("First name column", value="First Name")
        last_name_col = st.text_input("Last name column", value="Last Name")
        group_col = st.text_input("Group column", value="Group")
        email_col = st.text_input("Email column", value="Email")

        st.subheader("Assignment columns")
        scores_start_index = st.number_input(
            "Assignment columns start at index",
            min_value=0,
            value=4,
            step=1,
            help="0-based column index. Example: 4 means the 5th column onward are assignments.",
        )

        st.divider()
        st.markdown("**Output format per student:**")
        st.code("First Last, Group, Email")

    uploaded = st.file_uploader("Upload Excel file", type=["xlsx", "xls"])

    if not uploaded:
        st.info("Upload an Excel file to begin.")
        return

    try:
        df = pd.read_excel(uploaded)
    except Exception as e:
        st.error(f"Could not read Excel file: {e}")
        return

    st.subheader("Preview")
    st.dataframe(df, use_container_width=True)

    try:
        records = generate_missed_homework_report(
            df=df,
            first_name_col=first_name_col,
            last_name_col=last_name_col,
            group_col=group_col,
            email_col=email_col,
            scores_start_index=int(scores_start_index),
        )
    except Exception as e:
        st.error(str(e))
        return

    st.subheader("Detected students with missed homework")

    if not records:
        st.success("No students with missed homework (score = 0) were found.")
        # Still allow downloading a report that states none were found
        docx_bytes = build_word_report_bytes(records, title.strip() or "Missed Homework Report")
        st.download_button(
            "Download Word Report (.docx)",
            data=docx_bytes,
            file_name="missed_homework_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        return

    # Quick summary
    total_students = len(records)
    total_missed = sum(len(r["missed"]) for r in records)
    c1, c2 = st.columns(2)
    c1.metric("Students with misses", total_students)
    c2.metric("Total missed items", total_missed)

    # Table view
    table_rows = []
    for r in records:
        table_rows.append(
            {
                "Student": r["name"],
                "Group": r["group"],
                "Email": r["email"],
                "Missed assignments": ", ".join(r["missed"]),
            }
        )
    st.dataframe(pd.DataFrame(table_rows), use_container_width=True)

    st.divider()

    docx_bytes = build_word_report_bytes(records, title.strip() or "Missed Homework Report")
    st.download_button(
        "Download Word Report (.docx)",
        data=docx_bytes,
        file_name="missed_homework_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    if records:
        mail_df = build_mail_merge_dataframe(records)

        csv_bytes = mail_df.to_csv(
            index=False,
            quoting=csv.QUOTE_ALL  # Ensures line breaks are preserved
        ).encode("utf-8-sig")      # Outlook-friendly encoding (UTF-8 with BOM)

        st.download_button(
            label="Download Mail Merge CSV",
            data=csv_bytes,
            file_name="mail_merge_data.csv",
            mime="text/csv",
            use_container_width=True,
        )


if __name__ == "__main__":
    main()