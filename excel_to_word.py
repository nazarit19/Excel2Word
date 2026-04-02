import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document
from tkinter import ttk

def generate_missed_homework_report(file_path):
    df = pd.read_excel(file_path)
    score_columns = df.columns[4:]
    missed_assignments = {}

    for index, row in df.iterrows():
        student_name = f"{row['First Name']} {row['Last Name']}"
        missed_assignments[student_name] = []
        
        for column in score_columns:
            if row[column] == 0:
                missed_assignments[student_name].append(column)

    return missed_assignments

def write_report_to_word(missed_assignments, word_file_path):
    doc = Document()
    doc.add_heading('Missed Homework Report', 0)

    for student, assignments in missed_assignments.items():
        doc.add_heading(student, level=1)
        for assignment in assignments:
            doc.add_paragraph(f" {assignment}", style='ListBullet')
        doc.add_paragraph("")

    doc.save(word_file_path)

def browse_excel_file():
    filename = filedialog.askopenfilename(filetypes=(("Excel files", "*.xlsx;*.xls"), ("All files", "*.*")))
    excel_entry.delete(0, tk.END)
    excel_entry.insert(0, filename)

def browse_word_file():
    filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=(("Word Document", "*.docx"), ("All files", "*.*")))
    word_entry.delete(0, tk.END)
    word_entry.insert(0, filename)

def generate_report():
    excel_file_path = excel_entry.get()
    word_file_path = word_entry.get()

    if not excel_file_path or not word_file_path:
        messagebox.showwarning("Warning", "Please select both Excel and Word file paths.")
        return

    try:
        missed_assignments = generate_missed_homework_report(excel_file_path)
        write_report_to_word(missed_assignments, word_file_path)
        messagebox.showinfo("Success", "Report has been successfully generated.")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# Set up the GUI
root = tk.Tk()
root.title("Missed Homework Report Generator")

frame = ttk.Frame(root, padding="10")
frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

excel_label = ttk.Label(frame, text="Excel File:")
excel_label.grid(row=0, column=0, sticky=tk.W)
excel_entry = ttk.Entry(frame, width=50)
excel_entry.grid(row=0, column=1, sticky=(tk.W, tk.E))
browse_excel_btn = ttk.Button(frame, text="Browse...", command=browse_excel_file)
browse_excel_btn.grid(row=0, column=2, sticky=tk.W)

word_label = ttk.Label(frame, text="Save Word Report As:")
word_label.grid(row=1, column=0, sticky=tk.W)
word_entry = ttk.Entry(frame, width=50)
word_entry.grid(row=1, column=1, sticky=(tk.W, tk.E))
browse_word_btn = ttk.Button(frame, text="Browse...", command=browse_word_file)
browse_word_btn.grid(row=1, column=2, sticky=tk.W)

generate_btn = ttk.Button(frame, text="Generate Report", command=generate_report)
generate_btn.grid(row=2, column=0, columnspan=3)

root.mainloop()
